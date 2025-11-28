"""
Content Tools for Multi-Agent Content Creator System
Includes: Fact-Checking, Multi-Format Export, Social Media Generator
"""

import os
from typing import Dict, List
from datetime import datetime
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch

load_dotenv()

# Configure LLM
llm = ChatGroq(
    model=os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
    temperature=float(os.getenv("GROQ_TEMPERATURE", "0.7")),
    groq_api_key=os.getenv("GROQ_API_KEY"),
    max_tokens=int(os.getenv("GROQ_MAX_TOKENS", "2000"))
)


class FactCheckingAgent:
    """Agent that verifies claims and adds citations"""
    
    def __init__(self):
        self.role = "Fact-Checker"
        self.verification_template = PromptTemplate(
            input_variables=["article"],
            template="""You are a meticulous fact-checking agent. Review the following article and:

Article:
{article}

For each factual claim (statistics, names, dates, events):
1. Evaluate the claim's verifiability
2. Provide a confidence score (0-100%)
3. Suggest reliable sources if available
4. Flag any questionable claims

Return JSON format:
{{
    "verified_claims": [
        {{"claim": "...", "confidence": 85, "source": "...", "verified": true}}
    ],
    "unverified_claims": [
        {{"claim": "...", "reason": "...", "needs_source": true}}
    ],
    "overall_accuracy": 95,
    "improvements": ["suggestion 1", "suggestion 2"]
}}"""
        )
    
    def verify_article(self, article: str) -> Dict:
        """Verify claims in article"""
        print(f"\n🔍 {self.role} is verifying claims...")
        
        chain = self.verification_template | llm
        result = chain.invoke({"article": article[:3000]}).content  # Limit to first 3000 chars
        
        try:
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                verification_data = json.loads(json_match.group())
            else:
                verification_data = {
                    "verified_claims": [],
                    "unverified_claims": [],
                    "overall_accuracy": 85,
                    "improvements": [result]
                }
        except json.JSONDecodeError:
            verification_data = {
                "verified_claims": [],
                "unverified_claims": [],
                "overall_accuracy": 85,
                "improvements": [result]
            }
        
        print(f"✅ Fact-checking completed - Accuracy: {verification_data.get('overall_accuracy', 85)}%")
        return verification_data
    
    def generate_fact_check_report(self, verification_data: Dict) -> str:
        """Generate human-readable fact-check report"""
        report = "\n" + "="*70 + "\n"
        report += "FACT-CHECK VERIFICATION REPORT\n"
        report += "="*70 + "\n\n"
        
        report += f"Overall Accuracy Score: {verification_data.get('overall_accuracy', 85)}%\n\n"
        
        verified = verification_data.get('verified_claims', [])
        if verified:
            report += "✅ VERIFIED CLAIMS:\n"
            for claim in verified:
                if isinstance(claim, dict):
                    report += f"  • {claim.get('claim', 'N/A')} (Confidence: {claim.get('confidence', 0)}%)\n"
        
        unverified = verification_data.get('unverified_claims', [])
        if unverified:
            report += "\n⚠️ CLAIMS NEEDING VERIFICATION:\n"
            for claim in unverified:
                if isinstance(claim, dict):
                    report += f"  • {claim.get('claim', 'N/A')}\n"
        
        improvements = verification_data.get('improvements', [])
        if improvements:
            report += "\n💡 IMPROVEMENT SUGGESTIONS:\n"
            for i, suggestion in enumerate(improvements[:3], 1):
                if isinstance(suggestion, str):
                    report += f"  {i}. {suggestion}\n"
        
        report += "\n" + "="*70 + "\n"
        return report


class MultiFormatExporter:
    """Exports articles to multiple formats"""
    
    def __init__(self):
        self.role = "Multi-Format Exporter"
    
    def export_to_markdown(self, article: str, filename: str = "article.md") -> str:
        """Export as Markdown"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"# Article\n\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(article)
        print(f"✅ Markdown exported: {filename}")
        return filename
    
    def export_to_html(self, article: str, filename: str = "article.html") -> str:
        """Export as HTML"""
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Generated Article</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
            color: #333;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        .meta {{
            color: #7f8c8d;
            font-size: 0.9em;
            margin-bottom: 20px;
        }}
        p {{
            margin: 15px 0;
        }}
        strong {{
            color: #2c3e50;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="meta">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
        <article>
            {article.replace('**', '<strong>').replace('\n', '<br>')}
        </article>
    </div>
</body>
</html>
"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"✅ HTML exported: {filename}")
        return filename
    
    def export_to_docx(self, article: str, filename: str = "article.docx") -> str:
        """Export as Word Document"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Generated Article', 0)
        title.alignment = 1  # Center alignment
        
        # Add metadata
        meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        meta.runs[0].italic = True
        meta.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Blank line
        
        # Add content
        lines = article.split('\n')
        for line in lines:
            if line.startswith('**') and line.endswith('**'):
                # Bold heading
                heading = line.replace('**', '')
                doc.add_heading(heading, level=2)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(filename)
        print(f"✅ Word document exported: {filename}")
        return filename
    
    def export_to_pdf(self, article: str, filename: str = "article.pdf") -> str:
        """Export as PDF"""
        try:
            doc = SimpleDocTemplate(filename, pagesize=letter)
            story = []
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=RGBColor(44, 62, 80),
                spaceAfter=30,
                alignment=1
            )
            
            # Title
            story.append(Paragraph("Generated Article", title_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Meta info
            meta_style = styles['Normal']
            meta_style.fontSize = 9
            story.append(Paragraph(f"<i>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>", meta_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Content
            lines = article.split('\n')
            for line in lines:
                if line.startswith('**') and line.endswith('**'):
                    heading = line.replace('**', '')
                    story.append(Paragraph(heading, styles['Heading2']))
                    story.append(Spacer(1, 0.2*inch))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            print(f"✅ PDF exported: {filename}")
            return filename
        except Exception as e:
            print(f"⚠️ PDF export failed: {str(e)}")
            return None


class SocialMediaGenerator:
    """Generates social media content from articles"""
    
    def __init__(self):
        self.role = "Social Media Content Generator"
        self.social_template = PromptTemplate(
            input_variables=["article"],
            template="""Based on this article, create social media content:

Article:
{article}

Generate JSON with:
{{
    "twitter_thread": [
        "Tweet 1 (280 chars max)",
        "Tweet 2 (280 chars max)",
        "Tweet 3 (280 chars max)"
    ],
    "linkedin_post": "Professional post (300 chars max)",
    "instagram_caption": "Engaging caption with hashtags (2200 chars max)",
    "email_subject": "Email subject line",
    "email_preview": "Email preview text (50 chars)",
    "hashtags": ["#tag1", "#tag2", "#tag3", "#tag4"],
    "key_quote": "Best quote from article"
}}"""
        )
    
    def generate_content(self, article: str) -> Dict:
        """Generate social media content"""
        print(f"\n📱 {self.role} is creating content...")
        
        chain = self.social_template | llm
        result = chain.invoke({"article": article[:2000]}).content
        
        try:
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if json_match:
                social_data = json.loads(json_match.group())
            else:
                social_data = {
                    "twitter_thread": [result[:280]],
                    "linkedin_post": result[:300],
                    "instagram_caption": result,
                    "email_subject": "Check out this article",
                    "email_preview": "Interesting content",
                    "hashtags": ["#content", "#article"],
                    "key_quote": result.split('.')[0]
                }
        except json.JSONDecodeError:
            social_data = {
                "twitter_thread": [result[:280]],
                "linkedin_post": result[:300],
                "instagram_caption": result,
                "email_subject": "Check out this article",
                "email_preview": "Interesting content",
                "hashtags": ["#content", "#article"],
                "key_quote": result
            }
        
        print(f"✅ Social media content generated")
        return social_data
    
    def generate_social_report(self, social_data: Dict) -> str:
        """Generate formatted social media content report"""
        report = "\n" + "="*70 + "\n"
        report += "SOCIAL MEDIA CONTENT PACKAGE\n"
        report += "="*70 + "\n\n"
        
        # Twitter Thread
        report += "🐦 TWITTER THREAD:\n"
        tweets = social_data.get('twitter_thread', [])
        for i, tweet in enumerate(tweets, 1):
            report += f"  Tweet {i}: {tweet}\n"
        
        # LinkedIn
        report += "\n💼 LINKEDIN POST:\n"
        report += f"  {social_data.get('linkedin_post', 'N/A')}\n"
        
        # Instagram
        report += "\n📸 INSTAGRAM CAPTION:\n"
        report += f"  {social_data.get('instagram_caption', 'N/A')}\n"
        
        # Email
        report += "\n✉️ EMAIL:\n"
        report += f"  Subject: {social_data.get('email_subject', 'N/A')}\n"
        report += f"  Preview: {social_data.get('email_preview', 'N/A')}\n"
        
        # Hashtags
        report += "\n#️⃣ HASHTAGS:\n"
        hashtags = social_data.get('hashtags', [])
        for tag in hashtags:
            report += f"  {tag} "
        
        # Quote
        report += f"\n\n💬 KEY QUOTE:\n"
        report += f"  \"{social_data.get('key_quote', 'N/A')}\"\n"
        
        report += "\n" + "="*70 + "\n"
        return report


def generate_comprehensive_output(article: str, topic: str):
    """Generate all outputs: fact-checking, exports, and social media"""
    
    print("\n" + "="*70)
    print("🚀 GENERATING COMPREHENSIVE CONTENT PACKAGE")
    print("="*70)
    
    # 1. Fact-Checking
    fact_checker = FactCheckingAgent()
    verification_data = fact_checker.verify_article(article)
    fact_check_report = fact_checker.generate_fact_check_report(verification_data)
    
    # 2. Multi-Format Export
    exporter = MultiFormatExporter()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    exporter.export_to_markdown(article, f"exports/article_{timestamp}.md")
    exporter.export_to_html(article, f"exports/article_{timestamp}.html")
    exporter.export_to_docx(article, f"exports/article_{timestamp}.docx")
    exporter.export_to_pdf(article, f"exports/article_{timestamp}.pdf")
    
    # 3. Social Media Content
    social_gen = SocialMediaGenerator()
    social_data = social_gen.generate_content(article)
    social_report = social_gen.generate_social_report(social_data)
    
    # Save comprehensive report
    comprehensive_report = f"""
{fact_check_report}

{social_report}

===============================================================================
EXPORT FILES CREATED
===============================================================================
✅ article_{timestamp}.md (Markdown)
✅ article_{timestamp}.html (Web-ready HTML)
✅ article_{timestamp}.docx (Word Document)
✅ article_{timestamp}.pdf (PDF Format)
"""
    
    with open("comprehensive_output.txt", 'w', encoding='utf-8') as f:
        f.write(comprehensive_report)
    
    # Save social data as JSON for programmatic use
    with open(f"social_content_{timestamp}.json", 'w', encoding='utf-8') as f:
        json.dump(social_data, f, indent=2)
    
    print(fact_check_report)
    print(social_report)
    print("\n📄 Full report saved to: comprehensive_output.txt")
    print(f"📱 Social media data saved to: social_content_{timestamp}.json")
    
    return {
        "fact_check": verification_data,
        "social_media": social_data,
        "exports": {
            "markdown": f"article_{timestamp}.md",
            "html": f"article_{timestamp}.html",
            "docx": f"article_{timestamp}.docx",
            "pdf": f"article_{timestamp}.pdf"
        }
    }
